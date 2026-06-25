"""
cv_docx_generator.py — Word (.docx) CV Generator
Uses python-docx to produce a formatted, editable Word document.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io
import re


# ── Colors ───────────────────────────────────────────────────────────────────
BLACK  = RGBColor(0x11, 0x11, 0x11)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
LGRAY  = RGBColor(0xEE, 0xEE, 0xEE)
LINK   = RGBColor(0x00, 0x00, 0xCC)


def set_cell_bg(cell, color_hex: str):
    """Set background color of a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)


def set_cell_margins(cell, top=40, bottom=40, left=60, right=60):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:w'), str(val))
        el.set(qn('w:type'), 'dxa')
        tcMar.append(el)
    tcPr.append(tcMar)


def remove_cell_borders(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'none')
        tcBorders.append(b)
    tcPr.append(tcBorders)


def add_section_header(doc, title: str):
    """Full-width black row with white letter-spaced title."""
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    cell = table.cell(0, 0)
    set_cell_bg(cell, '111111')
    set_cell_margins(cell, top=50, bottom=50, left=100, right=100)

    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    spaced = ' '.join(title.upper())
    run = para.add_run(spaced)
    run.bold = True
    run.font.size = Pt(8)
    run.font.color.rgb = WHITE

    # Remove table border
    for row in table.rows:
        for c in row.cells:
            remove_cell_borders(c)

    doc.add_paragraph()  # small gap


def add_label_content_row(doc, label: str, content_lines: list,
                           right_text: str = '', sub_bullet: bool = False):
    """One row: left label column + right content column."""
    if not content_lines:
        return

    table = doc.add_table(rows=len(content_lines), cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Column widths: ~25% label, ~75% content
    label_w = Cm(4.5)
    content_w = Cm(13.0)

    for i, line in enumerate(content_lines):
        row = table.rows[i]
        label_cell = row.cells[0]
        content_cell = row.cells[1]

        set_cell_margins(label_cell, top=30, bottom=30, left=60, right=40)
        set_cell_margins(content_cell, top=30, bottom=30, left=40, right=60)

        label_cell.width = label_w
        content_cell.width = content_w

        # Label only on first row
        if i == 0 and label:
            lp = label_cell.paragraphs[0]
            lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
            lr = lp.add_run(label)
            lr.bold = True
            lr.font.size = Pt(7)

        # Content
        cp = content_cell.paragraphs[0]
        cp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        prefix = '– ' if sub_bullet else '■ '
        run = cp.add_run(prefix + line)
        run.font.size = Pt(7)

        for cell in [label_cell, content_cell]:
            remove_cell_borders(cell)

    doc.add_paragraph()


def generate_cv_docx(data: dict) -> bytes:
    """Generate Word CV from data dict. Returns bytes."""
    doc = Document()

    # ── Page setup ────────────────────────────────────────────────────────────
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(1.2)
    section.right_margin = Cm(1.2)
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.0)

    # Default paragraph style
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(7.5)

    # ── Header ─────────────────────────────────────────────────────────────────
    name = data.get('name', '').upper()
    gender = data.get('gender', '').upper()
    dob = data.get('dob', '')
    degree_line = data.get('degree_line', '').upper()

    header_table = doc.add_table(rows=1, cols=2)
    header_table.style = 'Table Grid'
    left_cell = header_table.cell(0, 0)
    right_cell = header_table.cell(0, 1)
    left_cell.width = Cm(10)
    right_cell.width = Cm(7.5)

    # Name
    lp = left_cell.paragraphs[0]
    lr = lp.add_run(' '.join(name))
    lr.bold = True
    lr.font.size = Pt(16)
    lr.font.color.rgb = BLACK

    # Right side: gender/dob + degree
    rp = right_cell.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    gd = f"{gender} , {dob}" if gender and dob else (gender or dob)
    if gd:
        rr = rp.add_run(gd)
        rr.font.size = Pt(8)

    if degree_line:
        rp2 = right_cell.add_paragraph()
        rp2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        rr2 = rp2.add_run(degree_line)
        rr2.font.size = Pt(8)

    for c in [left_cell, right_cell]:
        remove_cell_borders(c)
        set_cell_margins(c, top=60, bottom=60, left=60, right=60)

    # ── Spike points ───────────────────────────────────────────────────────────
    spikes = data.get('spike_points', [])
    if spikes:
        st_table = doc.add_table(rows=1, cols=1)
        sc = st_table.cell(0, 0)
        set_cell_bg(sc, '111111')
        set_cell_margins(sc, top=60, bottom=60, left=60, right=60)
        sp = sc.paragraphs[0]
        sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sr = sp.add_run('   |   '.join(s.strip() for s in spikes[:4]))
        sr.bold = True
        sr.font.size = Pt(8.5)
        sr.font.color.rgb = WHITE
        remove_cell_borders(sc)

    doc.add_paragraph()

    # ── Academic profile ───────────────────────────────────────────────────────
    ap = data.get('academic_profile', [])
    if ap:
        add_section_header(doc, 'ACADEMIC PROFILE')

        table = doc.add_table(rows=len(ap), cols=4)
        table.style = 'Table Grid'

        col_widths = [Cm(4), Cm(9.5), Cm(2.8), Cm(2.0)]
        for i, e in enumerate(ap):
            row = table.rows[i]
            vals = [e.get('degree',''), e.get('institution',''), e.get('score',''), e.get('year','')]
            for j, (cell, val) in enumerate(zip(row.cells, vals)):
                cell.width = col_widths[j]
                set_cell_margins(cell, top=40, bottom=40, left=60, right=60)
                if i % 2 == 0:
                    set_cell_bg(cell, 'F0F0F0')
                p = cell.paragraphs[0]
                r = p.add_run(val)
                r.font.size = Pt(7)

    doc.add_paragraph()

    # ── Work experience ───────────────────────────────────────────────────────
    we = data.get('work_experience', [])
    if we:
        add_section_header(doc, 'WORK EXPERIENCE')
        for exp in we:
            # Company + role header
            t = doc.add_table(rows=1, cols=2)
            t.style = 'Table Grid'
            lc = t.cell(0, 0)
            rc = t.cell(0, 1)
            lc.width = Cm(4.5)
            rc.width = Cm(13.0)
            for c in [lc, rc]:
                remove_cell_borders(c)
                set_cell_margins(c, top=40, bottom=20, left=60, right=60)

            lp = lc.paragraphs[0]
            lr = lp.add_run(exp.get('company', ''))
            lr.bold = True; lr.font.size = Pt(7.5)

            rp = rc.paragraphs[0]
            rp.alignment = WD_ALIGN_PARAGRAPH.LEFT
            rr = rp.add_run(exp.get('role', ''))
            rr.bold = True; rr.font.size = Pt(7.5)

            # Duration on right
            dur = exp.get('duration', '')
            if dur:
                dur_p = rp.add_run(f'    {dur}')
                dur_p.bold = True; dur_p.font.size = Pt(7)

            for key, label in [('responsibilities','Roles & Resp.'), ('initiatives','Initiatives'), ('achievements','Achievements')]:
                pts = exp.get(key, [])
                if pts:
                    add_label_content_row(doc, label, pts, sub_bullet=True)

    # ── Academic achievements ──────────────────────────────────────────────────
    aa = data.get('academic_achievements', {})
    has_aa = any(aa.get(k) for k in ('academic','competitions','scholarships'))
    if has_aa:
        add_section_header(doc, 'ACADEMIC ACHIEVEMENTS')
        for cat_key, cat_label in [('academic','Academic'),('competitions','Competitions'),('scholarships','Scholarships')]:
            pts = aa.get(cat_key, [])
            if not pts:
                continue
            lines = []
            for p in pts:
                if isinstance(p, dict):
                    t = p.get('text','')
                    y = str(p.get('year',''))
                    lines.append(f"{t}  {y}".strip())
                else:
                    lines.append(str(p))
            add_label_content_row(doc, cat_label, lines)

    # ── POR ────────────────────────────────────────────────────────────────────
    por = data.get('positions_of_responsibility', [])
    if por:
        add_section_header(doc, 'POSITIONS OF RESPONSIBILITY')
        for e in por:
            t = doc.add_table(rows=1, cols=2)
            t.style = 'Table Grid'
            lc = t.cell(0, 0); rc = t.cell(0, 1)
            lc.width = Cm(4.5); rc.width = Cm(13.0)
            for c in [lc, rc]:
                remove_cell_borders(c)
                set_cell_margins(c, top=40, bottom=20, left=60, right=60)
            lp = lc.paragraphs[0]
            lr = lp.add_run(e.get('organization',''))
            lr.bold = True; lr.font.size = Pt(7.5)
            rp = rc.paragraphs[0]
            role_txt = e.get('role','')
            yr = str(e.get('year',''))
            rr = rp.add_run(f"{role_txt}    {yr}" if yr else role_txt)
            rr.bold = True; rr.font.size = Pt(7.5)
            if e.get('bullets'):
                add_label_content_row(doc, '', e['bullets'], sub_bullet=True)

    # ── CIP ────────────────────────────────────────────────────────────────────
    cip = data.get('cip', {})
    all_cip = []
    for k in ('certifications','internships','projects'):
        all_cip.extend(cip.get(k,[]))
    if all_cip:
        add_section_header(doc, 'CERTIFICATIONS, INTERNSHIPS & PROJECTS')
        for e in all_cip:
            t = doc.add_table(rows=1, cols=2)
            t.style = 'Table Grid'
            lc = t.cell(0,0); rc = t.cell(0,1)
            lc.width = Cm(4.5); rc.width = Cm(13.0)
            for c in [lc, rc]:
                remove_cell_borders(c)
                set_cell_margins(c, top=40, bottom=20, left=60, right=60)
            lp = lc.paragraphs[0]
            lr = lp.add_run(e.get('organization',''))
            lr.bold = True; lr.font.size = Pt(7.5)
            rp = rc.paragraphs[0]
            title_txt = e.get('title','')
            dur = e.get('duration','')
            rr = rp.add_run(f"{title_txt}    {dur}" if dur else title_txt)
            rr.bold = True; rr.font.size = Pt(7.5)
            if e.get('bullets'):
                add_label_content_row(doc, '', e['bullets'], sub_bullet=True)

    # ── ECA ────────────────────────────────────────────────────────────────────
    eca = data.get('eca', {})
    if eca:
        add_section_header(doc, 'EXTRA CURRICULAR ACTIVITIES')
        ORDER = [
            'Debate/ Public Speaking','Sports / Adventure Sports','Management',
            'Cultural','Art & Design','Quizzing','Social Service','Technical','Literature','Others'
        ]
        extra = [k for k in eca if k not in ORDER]
        for cat in ORDER + extra:
            pts = eca.get(cat, [])
            if not pts: continue
            lines = []
            for p in pts:
                if isinstance(p, dict):
                    t = p.get('text','')
                    y = str(p.get('year',''))
                    lines.append(f"{t}  {y}".strip())
                else:
                    lines.append(str(p))
            add_label_content_row(doc, cat, lines)

    # ── Footer ─────────────────────────────────────────────────────────────────
    parts = [p for p in [data.get('linkedin'), data.get('email'), data.get('phone')] if p]
    if parts:
        fp = doc.add_paragraph('   |   '.join(parts))
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in fp.runs:
            run.font.size = Pt(7)
            run.font.color.rgb = LINK

    # ── Output as bytes ────────────────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
