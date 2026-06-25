"""
CV Standardization Agent — Streamlit Web App
Upload a CV → AI parses it → fills gaps with Q&A → outputs PDF + Word
"""

import streamlit as st
import requests
import json
import re
import os
import sys
import tempfile
from pathlib import Path
import io

# ── Page config ─────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="CV Standardization Agent",
    page_icon="📄",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ── Inject CSS ───────────────────────────────────────────────────────────────
st.markdown("""
<style>
  .section-bar {
    background:#111; color:#fff; padding:6px 14px;
    font-weight:700; letter-spacing:2px; margin:14px 0 6px;
    font-size:13px;
  }
  .spike-bar {
    background:#111; color:#fff; padding:8px; text-align:center;
    font-weight:700; margin:8px 0 12px; font-size:13px; letter-spacing:1px;
  }
  .missing-field { border:2px solid #ff4b4b !important; }
  .info-box {
    background:#f0f4ff; border-left:4px solid #4a6cf7;
    padding:10px 14px; margin:8px 0; border-radius:4px;
  }
  h1 { font-size:2rem !important; }
</style>
""", unsafe_allow_html=True)

# ── Constants ────────────────────────────────────────────────────────────────

PARSE_PROMPT = """You are a CV parsing expert. Extract ALL information from this CV text and return it as valid JSON.

Apply these formatting rules while extracting:
- Numbers ≥100,000 → mn/bn (150000 → 0.15 mn, 1200000 → 1.2 mn)
- Numbers 1,000–99,999 → add commas (3000 → 3,000)
- Currency symbols → INR/USD/GBP/EUR (₹ → INR, $ → USD)
- Percentile: use %ile (not % or percentile)
- Use capital Top for ranks: Top 5 %ile, Top 3 ranks
- Work experience dates: Jun'21-Aug'22 format
- Other section years: just the ending year e.g. 2021
- Degree abbreviations: B.Sc.(H), B.Tech., B.Com.(H), MBA, M.Tech., CA, etc.
- School format: "School Name, City (BOARD)" e.g. "DPS, Ranchi (CBSE)"
- All bullets must start with action verbs
- Remove Pvt Ltd / Private Limited from company names
- Single quotes only, no double quotes in content

Return ONLY this JSON schema (no explanation):
{
  "name": "FirstName LastName",
  "gender": "",
  "dob": "YYYY-MM-DD",
  "degree_line": "e.g. B.Tech CSE | 2018-22",
  "spike_points": [],
  "academic_profile": [
    {"degree": "", "institution": "", "score": "", "year": ""}
  ],
  "work_experience": [
    {"company": "", "role": "", "duration": "", "responsibilities": [], "initiatives": [], "achievements": []}
  ],
  "academic_achievements": {
    "academic": [{"text": "", "year": ""}],
    "competitions": [{"text": "", "year": ""}],
    "scholarships": [{"text": "", "year": ""}]
  },
  "positions_of_responsibility": [
    {"organization": "", "role": "", "year": "", "bullets": []}
  ],
  "cip": {
    "certifications": [{"organization": "", "title": "", "duration": "", "bullets": []}],
    "internships": [{"organization": "", "title": "", "duration": "", "bullets": []}],
    "projects": [{"organization": "", "title": "", "duration": "", "bullets": []}]
  },
  "eca": {
    "Sports / Adventure Sports": [{"text": "", "year": ""}],
    "Others": [{"text": "Hobbies: ...", "year": ""}]
  },
  "linkedin": "",
  "email": "",
  "phone": ""
}

ECA category names (use exactly): Debate/ Public Speaking | Sports / Adventure Sports | Management | Cultural | Art & Design | Quizzing | Social Service | Technical | Literature | Others

CV TEXT:
{cv_text}"""

IMPROVE_PROMPT = """You are a CV writing expert. Given this CV bullet point, improve it to:
1. Start with a strong action verb (from: Developed, Led, Built, Drove, Achieved, Launched, Managed, Designed, Reduced, Increased, etc.)
2. Include metrics/numbers wherever possible
3. Keep it concise — single line
4. Apply formatting: numbers ≥100,000 → mn/bn, use INR/USD not symbols

Original bullet: {bullet}
Context: {context}

Return ONLY the improved single-line bullet. No explanation."""

# ── Helpers ──────────────────────────────────────────────────────────────────

def get_api_key() -> str:
    """Get API key from env var, Streamlit secrets, or session state."""
    # Render / any host: use environment variable
    import os
    env_key = os.environ.get("GROQ_API_KEY", "")
    if env_key:
        return env_key
    # Streamlit Cloud: use secrets
    try:
        if "GEMINI_API_KEY" in st.secrets:
            return st.secrets["GEMINI_API_KEY"]
    except Exception:
        pass
    return st.session_state.get("api_key", "")


def extract_text(uploaded_file) -> str:
    """Extract text from uploaded PDF or DOCX."""
    suffix = Path(uploaded_file.name).suffix.lower()
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as f:
        f.write(uploaded_file.read())
        tmp_path = f.name

    try:
        if suffix == ".pdf":
            import pdfplumber
            with pdfplumber.open(tmp_path) as pdf:
                return "\n".join(p.extract_text() or "" for p in pdf.pages)
        elif suffix in (".docx", ".doc"):
            from docx import Document
            doc = Document(tmp_path)
            lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells if c.text.strip()]
                    if cells:
                        lines.append(" | ".join(cells))
            return "\n".join(lines)
        elif suffix in (".txt", ".md"):
            return uploaded_file.read().decode("utf-8", errors="ignore")
        else:
            return ""
    finally:
        os.unlink(tmp_path)


def parse_cv_with_gemini(raw_text: str, api_key: str) -> dict:
    """Send raw CV text to Groq API and get structured JSON back."""
    prompt = PARSE_PROMPT.replace("{cv_text}", raw_text[:12000])
    url = "https://api.groq.com/openai/v1/chat/completions"
    headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}
    payload = {
        "model": "llama-3.3-70b-versatile",
        "messages": [{"role": "user", "content": prompt}],
        "max_tokens": 4096,
        "temperature": 0.1,
    }
    resp = requests.post(url, json=payload, headers=headers, timeout=60)
    resp.raise_for_status()
    content = resp.json()["choices"][0]["message"]["content"]

    # Extract JSON block
    match = re.search(r"\{[\s\S]*\}", content)
    if match:
        try:
            return json.loads(match.group())
        except json.JSONDecodeError:
            pass
    return {}


def find_missing_fields(data: dict) -> list[tuple[str, str]]:
    """Return list of (field_key, question) for missing required fields."""
    gaps = []
    if not data.get("name") or len(data["name"].split()) > 2:
        gaps.append(("name", "Full name (FirstName LastName — max 2 parts):"))
    if not data.get("gender"):
        gaps.append(("gender", "Gender (Female / Male / Non-binary / Prefer not to say):"))
    if not data.get("dob"):
        gaps.append(("dob", "Date of birth (YYYY-MM-DD, e.g. 2001-09-18):"))

    spikes = data.get("spike_points", [])
    if len(spikes) < 4:
        for i in range(len(spikes), 4):
            gaps.append((f"spike_{i}", f"Spike Point {i+1} (1–3 words, Pascal Case, e.g. 'National Rank 2'):"))

    if not data.get("academic_profile"):
        gaps.append(("ap_note", "⚠ No academic qualifications found — please add them below."))

    eca = data.get("eca", {})
    others = eca.get("Others", [])
    has_hobbies = any("hobbies" in str(p).lower() for p in others)
    if not has_hobbies:
        gaps.append(("hobbies", "Hobbies (required as last line, e.g. 'Playing Volleyball, Reading, Chess'):"))

    if not data.get("linkedin") and not data.get("email"):
        gaps.append(("linkedin", "LinkedIn profile URL (https://linkedin.com/in/yourhandle):"))
        gaps.append(("email", "Email address:"))

    return gaps


def generate_pdf(data: dict) -> bytes:
    """Generate standardized CV as PDF bytes."""
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as f:
        tmp_path = f.name

    # Import here to avoid circular issues
    sys.path.insert(0, os.path.dirname(__file__))
    from cv_generator import CV

    cv = CV(data, tmp_path)
    cv.generate()

    with open(tmp_path, "rb") as f:
        pdf_bytes = f.read()
    os.unlink(tmp_path)
    return pdf_bytes


def generate_docx(data: dict) -> bytes:
    """Generate standardized CV as Word .docx bytes."""
    sys.path.insert(0, os.path.dirname(__file__))
    from cv_docx_generator import generate_cv_docx
    return generate_cv_docx(data)


# ── Session state init ───────────────────────────────────────────────────────

def init_state():
    defaults = {
        "stage": "upload",      # upload → chat → download
        "raw_text": "",
        "cv_data": {},
        "api_key": "",
        "filename": "",
        "chat_messages": [],    # [{role, content}]
        "chat_queue": [],       # [(field_key, question)]
        "chat_done": False,
        "spike_suggestions": [],
        "awaiting_spike_choice": False,
    }
    for k, v in defaults.items():
        if k not in st.session_state:
            st.session_state[k] = v


# ── CHAT HELPERS ─────────────────────────────────────────────────────────────

def suggest_spike_points(data: dict, api_key: str) -> list:
    """Ask AI to suggest 4 spike points based on parsed CV."""
    summary = json.dumps({k: v for k, v in data.items() if k != "eca"}, indent=2)[:3000]
    prompt = f"""Based on this CV, suggest exactly 4 spike points for the header bar.

Rules:
- 2-4 words each, Pascal Case (e.g. "National Rank 2", "Ex-Deloitte Analyst")
- Cover 4 different areas: academics, work/internships, positions of responsibility, ECA or skills
- Be specific — use company names, ranks, numbers where possible
- These are the person's top 4 career highlights at a glance

CV data:
{summary}

Return ONLY a JSON array of exactly 4 strings. No explanation.
Example: ["IIT JEE Rank 245", "Ex-McKinsey Intern", "Student Council VP", "State-Level Swimmer"]"""

    try:
        url = "https://api.groq.com/openai/v1/chat/completions"
        headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}
        payload = {
            "model": "llama-3.3-70b-versatile",
            "messages": [{"role": "user", "content": prompt}],
            "max_tokens": 200,
            "temperature": 0.3,
        }
        resp = requests.post(url, json=payload, headers=headers, timeout=30)
        resp.raise_for_status()
        text = resp.json()["choices"][0]["message"]["content"]
        match = re.search(r'\[.*?\]', text, re.DOTALL)
        if match:
            return json.loads(match.group())
    except Exception:
        pass
    return []


def build_chat_queue(data: dict) -> list:
    """Return ordered list of (field_key, question) covering ALL CV sections."""
    q = []

    # ── Name — ALWAYS confirm (parsing often gets it wrong) ───────────────────
    parsed_name = data.get("name", "")
    if parsed_name:
        q.append(("name", f"I read your name as **{parsed_name}**. Is that correct?\nType **'yes'** to confirm, or type your correct full name."))
    else:
        q.append(("name", "What is your **full name**? (e.g. Arya Narayani — first and last name only)"))

    # ── Other basic info ──────────────────────────────────────────────────────
    if not data.get("gender"):
        q.append(("gender", "What is your **gender**? (Male / Female / Non-binary / Prefer not to say)"))
    if not data.get("dob"):
        q.append(("dob", "What is your **date of birth**? (format: YYYY-MM-DD, e.g. 1998-09-27)"))
    if not data.get("degree_line"):
        q.append(("degree_line", "What should appear as your **degree line** under your name? (e.g. MBA-HR | 2021-23  or  B.Tech CSE | 2018-22)"))

    # ── Academic profile ──────────────────────────────────────────────────────
    if not data.get("academic_profile"):
        q.append(("ap_note", "I couldn't find your qualifications. Please list them, one per line:\n`Degree | Institution | Score | Year`\ne.g.:\n`MBA-HR | Sri Balaji University, Pune | 7.8/10 | 2021-23`\n`BBA | ISBM, Pune (University of Pune) | 68% | 2017-20`"))

    # ── Spike points ──────────────────────────────────────────────────────────
    q.append(("spike_points", "SPIKE_GENERATION"))

    # ── Academic achievements — always ask ────────────────────────────────────
    aa = data.get("academic_achievements", {})
    aa_count = len(aa.get("academic", []) + aa.get("competitions", []) + aa.get("scholarships", []))
    q.append(("aa_academic", (
        f"I found {aa_count} academic achievement(s) in your CV. " if aa_count else "I didn't find any academic achievements. "
    ) + "Do you have any **academic prizes, ranks, or merit certificates**? "
      "List them one per line with year at end, or type **'none'** to skip.\n\n"
      "_e.g._ `Ranked Top 10 in State in Class XII CBSE | 2020`"))

    q.append(("aa_competitions", (
        "Any **competitions, case studies, hackathons, or B-school events** you've won or participated in? "
        "One per line with year, or type **'none'** to skip.\n\n"
        "_e.g._ `National Finalist at XYZ Case Competition among 4,000+ participants | 2023`"
    )))

    q.append(("aa_scholarships", (
        "Any **scholarships or financial awards**? One per line with year, or type **'none'** to skip.\n\n"
        "_e.g._ `Awarded INR 0.5 mn by Mirae Asset Foundation for holistic excellence | 2024`"
    )))

    # ── POR — always ask ──────────────────────────────────────────────────────
    por_count = len(data.get("positions_of_responsibility", []))
    q.append(("por_extra", (
        f"I found {por_count} position(s) of responsibility. " if por_count else "I didn't find any positions of responsibility. "
    ) + "Any **additional roles** — college clubs, societies, student council, sports captain, NGO, etc.? "
      "One per line in format: `Organization | Role | Year`, or type **'none'** to skip.\n\n"
      "_e.g._ `Ramjas College | Vice President, Zoological Society | 2022`"))

    # ── CIP — always ask ─────────────────────────────────────────────────────
    cip = data.get("cip", {})
    cip_count = sum(len(cip.get(k, [])) for k in ("certifications", "internships", "projects"))
    q.append(("cip_extra", (
        f"I found {cip_count} certification/internship/project(s). " if cip_count else "I didn't find any certifications, internships, or projects. "
    ) + "Any **additional ones** to add? One per line: `Organization | Title | Duration`, or type **'none'** to skip.\n\n"
      "_e.g._ `Coursera | Google Data Analytics Certificate | 2023`\n"
      "_e.g._ `Self | E-Commerce Startup — built and sold 3,000+ units | 2022`"))

    # ── ECA — always ask ─────────────────────────────────────────────────────
    eca = data.get("eca", {})
    eca_count = sum(len(v) for v in eca.values())
    q.append(("eca_extra", (
        f"I found {eca_count} extracurricular activity(ies). " if eca_count else "I didn't find any extracurriculars. "
    ) + "Any **sports, debate, cultural, social service, literature, or technical activities**? "
      "One per line: `Category | Achievement | Year`, or type **'none'** to skip.\n\n"
      "_e.g._ `Sports | Secured Gold in Volleyball at Inter-College Tournament | 2022`\n"
      "_e.g._ `Social Service | Raised INR 50,000 for 70+ underprivileged children via Enactus | 2021`"))

    # ── Hobbies ───────────────────────────────────────────────────────────────
    has_hobbies = any("hobbies" in str(p).lower() for p in eca.get("Others", []))
    if not has_hobbies:
        q.append(("hobbies", "What are your **hobbies**? (Last line of CV, e.g. Watching Anime, Reading Fiction, Playing Volleyball)"))

    # ── Contact — ALWAYS ask to confirm or add ───────────────────────────────
    parsed_li = data.get("linkedin", "")
    if parsed_li:
        q.append(("linkedin", f"I found your LinkedIn as **{parsed_li}**. Type **'yes'** to keep it, or paste the correct URL."))
    else:
        q.append(("linkedin", "What is your **LinkedIn URL**? (e.g. https://linkedin.com/in/yourhandle) — type **'skip'** to leave blank"))

    parsed_email = data.get("email", "")
    if parsed_email:
        q.append(("email", f"I found your email as **{parsed_email}**. Type **'yes'** to keep it, or type the correct email."))
    else:
        q.append(("email", "What is your **email address**? — type **'skip'** to leave blank"))

    return q


_SKIP_WORDS = {"skip", "none", "no", "n/a", "na", "nil", "-", "nope", "nothing"}

def apply_chat_answer(field_key: str, answer: str, data: dict) -> dict:
    """Update cv_data based on user's chat answer."""
    answer = answer.strip()
    answer_lower = answer.lower()

    # Global skip for all content fields
    if answer_lower in _SKIP_WORDS and field_key not in ("name", "linkedin", "email", "gender", "dob", "degree_line"):
        return data

    if field_key == "name":
        # "yes" / "correct" / "right" = keep parsed name
        if answer_lower not in ("yes", "y", "correct", "right", "ok", "okay", "yep", "yup"):
            data["name"] = answer
    elif field_key == "gender":
        data["gender"] = answer
    elif field_key == "dob":
        data["dob"] = answer
    elif field_key == "degree_line":
        data["degree_line"] = answer
    elif field_key == "ap_note":
        entries = []
        for line in answer.splitlines():
            parts = [p.strip() for p in line.split("|")]
            if len(parts) >= 2:
                entries.append({
                    "degree": parts[0] if len(parts) > 0 else "",
                    "institution": parts[1] if len(parts) > 1 else "",
                    "score": parts[2] if len(parts) > 2 else "",
                    "year": parts[3] if len(parts) > 3 else "",
                })
        if entries:
            data["academic_profile"] = entries
    elif field_key == "spike_points":
        spikes = [s.strip() for s in answer.splitlines() if s.strip()]
        data["spike_points"] = spikes[:4]

    elif field_key == "aa_academic":
        lines = [l.strip() for l in answer.splitlines() if l.strip() and l.strip().lower() not in _SKIP_WORDS]
        if lines:
            aa = data.setdefault("academic_achievements", {})
            existing = aa.get("academic", [])
            for line in lines:
                parts = [p.strip() for p in line.rsplit("|", 1)]
                text = parts[0]
                year = parts[1] if len(parts) > 1 else ""
                existing.append({"text": text, "year": year})
            aa["academic"] = existing

    elif field_key == "aa_competitions":
        lines = [l.strip() for l in answer.splitlines() if l.strip() and l.strip().lower() not in _SKIP_WORDS]
        if lines:
            aa = data.setdefault("academic_achievements", {})
            existing = aa.get("competitions", [])
            for line in lines:
                parts = [p.strip() for p in line.rsplit("|", 1)]
                text = parts[0]
                year = parts[1] if len(parts) > 1 else ""
                existing.append({"text": text, "year": year})
            aa["competitions"] = existing

    elif field_key == "aa_scholarships":
        lines = [l.strip() for l in answer.splitlines() if l.strip() and l.strip().lower() not in _SKIP_WORDS]
        if lines:
            aa = data.setdefault("academic_achievements", {})
            existing = aa.get("scholarships", [])
            for line in lines:
                parts = [p.strip() for p in line.rsplit("|", 1)]
                text = parts[0]
                year = parts[1] if len(parts) > 1 else ""
                existing.append({"text": text, "year": year})
            aa["scholarships"] = existing

    elif field_key == "por_extra":
        lines = [l.strip() for l in answer.splitlines() if l.strip() and l.strip().lower() not in _SKIP_WORDS]
        if lines:
            por = data.setdefault("positions_of_responsibility", [])
            for line in lines:
                parts = [p.strip() for p in line.split("|")]
                por.append({
                    "organization": parts[0] if len(parts) > 0 else "",
                    "role": parts[1] if len(parts) > 1 else "",
                    "year": parts[2] if len(parts) > 2 else "",
                    "bullets": [],
                })

    elif field_key == "cip_extra":
        lines = [l.strip() for l in answer.splitlines() if l.strip() and l.strip().lower() not in _SKIP_WORDS]
        if lines:
            cip = data.setdefault("cip", {})
            projects = cip.get("projects", [])
            for line in lines:
                parts = [p.strip() for p in line.split("|")]
                projects.append({
                    "organization": parts[0] if len(parts) > 0 else "",
                    "title": parts[1] if len(parts) > 1 else line,
                    "duration": parts[2] if len(parts) > 2 else "",
                    "bullets": [],
                })
            cip["projects"] = projects

    elif field_key == "eca_extra":
        lines = [l.strip() for l in answer.splitlines() if l.strip() and l.strip().lower() not in _SKIP_WORDS]
        if lines:
            eca = data.setdefault("eca", {})
            for line in lines:
                parts = [p.strip() for p in line.split("|")]
                cat = parts[0] if len(parts) > 0 else "Others"
                text = parts[1] if len(parts) > 1 else line
                year = parts[2] if len(parts) > 2 else ""
                bucket = eca.setdefault(cat, [])
                bucket.append({"text": text, "year": year})

    elif field_key == "hobbies":
        eca = data.get("eca", {})
        eca["Others"] = [{"text": f"Hobbies: {answer}", "year": ""}]
        data["eca"] = eca
    elif field_key == "linkedin":
        if answer_lower in _SKIP_WORDS:
            data.pop("linkedin", None)
        elif answer_lower not in ("yes", "y", "correct", "ok", "okay"):
            data["linkedin"] = answer
        # "yes" = keep existing parsed value
    elif field_key == "email":
        if answer_lower in _SKIP_WORDS:
            data.pop("email", None)
        elif answer_lower not in ("yes", "y", "correct", "ok", "okay"):
            data["email"] = answer
        # "yes" = keep existing parsed value
    return data


# ── SIDEBAR ──────────────────────────────────────────────────────────────────

def sidebar():
    with st.sidebar:
        st.title("⚙️ Settings")

        # API key (only show if not already configured via env/secrets)
        import os
        _has_key = bool(os.environ.get("GEMINI_API_KEY")) or bool(get_api_key())
        if not _has_key:
            st.markdown("**Groq API Key**")
            key = st.text_input(
                "Enter your Groq API key",
                type="password",
                value=st.session_state.api_key,
                placeholder="gsk_...",
                help="Get a free key at console.groq.com",
            )
            st.session_state.api_key = key
            if not key:
                st.warning("🔑 API key required to process CVs")
                st.markdown("[Get a free API key →](https://console.groq.com/)")
        else:
            st.success("✅ API key configured")

        st.divider()
        st.markdown("**How it works**")
        st.markdown("""
1. 📤 Upload your CV (PDF or Word)
2. 🤖 AI extracts all details
3. ✏️ Fill in any missing info
4. 📥 Download your standardized CV
        """)

        st.divider()
        st.markdown("**Output format**")
        st.markdown("✅ PDF\n✅ Word (.docx)")

        if st.session_state.stage != "upload":
            st.divider()
            if st.button("🔄 Start Over", use_container_width=True):
                for k in ["stage", "raw_text", "cv_data", "filename", "chat_messages",
                          "chat_queue", "chat_done", "spike_suggestions", "awaiting_spike_choice"]:
                    st.session_state.pop(k, None)
                st.session_state.stage = "upload"
                st.rerun()


# ── STAGE 1: UPLOAD ──────────────────────────────────────────────────────────

def stage_upload():
    st.title("📄 CV Standardization Agent")
    st.markdown("Upload your CV and get a clean, standardized version — correctly formatted, properly structured, ready to submit.")

    col1, col2 = st.columns([2, 1])
    with col1:
        uploaded = st.file_uploader(
            "Upload your CV",
            type=["pdf", "docx", "doc", "txt"],
            help="Accepted: PDF, Word (.docx), or plain text",
        )

        if uploaded:
            st.session_state.filename = uploaded.name
            api_key = get_api_key()

            if not api_key:
                st.error("Please enter your Google Gemini API key in the sidebar first.")
                return

            if st.button("🚀 Process CV", type="primary", use_container_width=True):
                with st.spinner("Extracting text from your CV..."):
                    raw = extract_text(uploaded)

                if not raw.strip():
                    st.error("Could not extract text from this file. Please try a PDF or DOCX.")
                    return

                st.session_state.raw_text = raw

                with st.spinner("AI is reading and structuring your CV..."):
                    try:
                        parsed = parse_cv_with_gemini(raw, api_key)
                    except Exception as e:
                        if "API_KEY" in str(e) or "invalid" in str(e).lower():
                            st.error("❌ Invalid API key. Please check your key in the sidebar.")
                        else:
                            st.error(f"❌ Error: {e}")
                        return

                if not parsed:
                    st.error("Could not parse the CV. Please ensure it has readable text.")
                    return

                st.session_state.cv_data = parsed
                st.session_state.stage = "chat"
                st.rerun()

    with col2:
        st.markdown("""
<div class="info-box">
<b>What gets standardized:</b><br>
• Numbers → 0.5 mn, 3,000<br>
• Currency → INR, USD<br>
• Percentile → %ile<br>
• Dates → Jun'22, 2022-23<br>
• Degrees → B.Sc.(H), B.Tech.<br>
• Schools → Name, City (Board)<br>
• Bullets → action verbs<br>
• No blank gaps on the page
</div>
        """, unsafe_allow_html=True)

    st.divider()
    st.markdown("#### Or start from scratch")
    if st.button("✍️ Build CV from scratch (no upload)", use_container_width=False):
        st.session_state.cv_data = {
            "name": "", "gender": "", "dob": "", "degree_line": "",
            "spike_points": [], "academic_profile": [], "work_experience": [],
            "academic_achievements": {"academic": [], "competitions": [], "scholarships": []},
            "positions_of_responsibility": [], "cip": {"certifications": [], "internships": [], "projects": []},
            "eca": {}, "linkedin": "", "email": "", "phone": ""
        }
        st.session_state.stage = "chat"
        st.rerun()


# ── STAGE 2: CHAT ─────────────────────────────────────────────────────────────

def stage_chat():
    st.title("💬 Let's build your CV")
    if st.session_state.filename:
        st.caption(f"Source: {st.session_state.filename}")

    data = st.session_state.cv_data
    api_key = get_api_key()

    # ── Initialize chat on first entry ───────────────────────────────────────
    if not st.session_state.chat_messages:
        # Build welcome summary
        name = data.get("name", "")
        we_count = len(data.get("work_experience", []))
        ap_count = len(data.get("academic_profile", []))
        lines = [f"I've read your CV{' for ' + name if name else ''}. Here's what I found:"]
        if ap_count: lines.append(f"✅ {ap_count} academic qualification(s)")
        if we_count: lines.append(f"✅ {we_count} work experience(s)")
        if data.get("linkedin"): lines.append(f"✅ LinkedIn profile")
        lines.append("\nI'll ask you a few questions to fill in any gaps. Let's go!")

        st.session_state.chat_messages = [{"role": "assistant", "content": "\n".join(lines)}]
        queue = build_chat_queue(data)
        st.session_state.chat_queue = queue

        # Ask first question
        if queue:
            first_key, first_q = queue[0]
            if first_q == "SPIKE_GENERATION":
                _ask_spike_question(data, api_key)
            else:
                st.session_state.chat_messages.append({"role": "assistant", "content": first_q})
        else:
            st.session_state.chat_done = True
            st.session_state.chat_messages.append({"role": "assistant", "content": "✅ Everything looks complete! Click **Generate CV** below."})

    # ── Render all messages ───────────────────────────────────────────────────
    for msg in st.session_state.chat_messages:
        with st.chat_message(msg["role"]):
            st.markdown(msg["content"])

    # ── Generate button when done ─────────────────────────────────────────────
    if st.session_state.chat_done:
        st.divider()
        if st.button("✅ Generate My CV", type="primary", use_container_width=True):
            st.session_state.stage = "download"
            st.rerun()
        return

    # ── Chat input ────────────────────────────────────────────────────────────
    user_input = st.chat_input("Type your answer here...")
    if not user_input:
        return

    # Show user message
    st.session_state.chat_messages.append({"role": "user", "content": user_input})

    # Process answer
    queue = st.session_state.chat_queue
    if not queue:
        st.session_state.chat_done = True
        st.rerun()
        return

    field_key, question = queue[0]

    # Handle spike points choice
    if field_key == "spike_points" and st.session_state.awaiting_spike_choice:
        if user_input.strip().lower() in ("use these", "yes", "use", "accept", "ok", "okay"):
            data["spike_points"] = st.session_state.spike_suggestions
            st.session_state.chat_messages.append({
                "role": "assistant",
                "content": f"✅ Got it! Using: **{' | '.join(data['spike_points'])}**"
            })
        else:
            # User typed their own spikes
            spikes = [s.strip() for s in user_input.splitlines() if s.strip()]
            if len(spikes) < 4:
                # Try comma-separated too
                spikes = [s.strip() for s in user_input.split(",") if s.strip()]
            data["spike_points"] = spikes[:4]
            st.session_state.chat_messages.append({
                "role": "assistant",
                "content": f"✅ Got it! Spike points: **{' | '.join(data['spike_points'])}**"
            })
        st.session_state.awaiting_spike_choice = False
        queue.pop(0)
    else:
        # Normal answer
        data = apply_chat_answer(field_key, user_input, data)
        queue.pop(0)
        st.session_state.chat_messages.append({"role": "assistant", "content": "✅ Got it!"})

    st.session_state.cv_data = data
    st.session_state.chat_queue = queue

    # Ask next question
    if queue:
        next_key, next_q = queue[0]
        if next_q == "SPIKE_GENERATION":
            _ask_spike_question(data, api_key)
        else:
            st.session_state.chat_messages.append({"role": "assistant", "content": next_q})
    else:
        st.session_state.chat_done = True
        st.session_state.chat_messages.append({
            "role": "assistant",
            "content": "✅ All done! Click **Generate CV** below to get your standardized PDF and Word file."
        })

    st.rerun()


def _ask_spike_question(data: dict, api_key: str):
    """Generate spike suggestions and ask the spike question."""
    with st.spinner("Analyzing your CV for spike points..."):
        suggestions = suggest_spike_points(data, api_key)
    st.session_state.spike_suggestions = suggestions

    if suggestions:
        msg = (
            "Now for your **Spike Points** — 4 highlights shown in the header bar of your CV.\n\n"
            "Based on your CV, here's what I'd suggest:\n\n"
            + "\n".join(f"{i+1}. **{s}**" for i, s in enumerate(suggestions))
            + "\n\n---\n"
            "Type **'use these'** to accept them, or type your own 4 spike points (one per line).\n\n"
            "💡 **Tip:** Manually written spike points almost always work better — "
            "only you know which role you're targeting and what to highlight for it."
        )
        st.session_state.awaiting_spike_choice = True
    else:
        msg = (
            "Now for your **Spike Points** — 4 highlights shown in the header bar of your CV.\n\n"
            "These should be your 4 biggest achievements in 2-4 words each (Pascal Case).\n"
            "Cover: Academics | Work | Positions | ECA/Skills\n\n"
            "Example:\n`National Rank 2\nEx-Deloitte Analyst\nStudent Council VP\nState-Level Swimmer`\n\n"
            "💡 Type one per line."
        )
        st.session_state.awaiting_spike_choice = False
        # Update queue so it just accepts the typed answer directly
    st.session_state.chat_messages.append({"role": "assistant", "content": msg})


# ── STAGE 3: DOWNLOAD ─────────────────────────────────────────────────────────

def stage_download():
    st.title("🎉 Your Standardized CV is Ready!")

    data = st.session_state.cv_data
    name = data.get("name", "CV")

    col1, col2 = st.columns(2)

    with col1:
        st.markdown("### 📥 Download")

        # PDF
        with st.spinner("Generating PDF..."):
            try:
                pdf_bytes = generate_pdf(data)
                st.download_button(
                    label="⬇️ Download PDF",
                    data=pdf_bytes,
                    file_name=f"{name.replace(' ', '_')}_CV.pdf",
                    mime="application/pdf",
                    use_container_width=True,
                    type="primary",
                )
            except Exception as e:
                st.error(f"PDF error: {e}")

        st.markdown("")

        # Word
        with st.spinner("Generating Word document..."):
            try:
                docx_bytes = generate_docx(data)
                st.download_button(
                    label="⬇️ Download Word (.docx)",
                    data=docx_bytes,
                    file_name=f"{name.replace(' ', '_')}_CV.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                )
            except Exception as e:
                st.error(f"Word error: {e}")

        st.markdown("")
        if st.button("✏️ Edit CV", use_container_width=True):
            st.session_state.stage = "chat"
            st.rerun()

    with col2:
        st.markdown("### 📋 CV Summary")
        spikes = data.get("spike_points", [])
        if spikes:
            st.markdown(f"""<div class="spike-bar">{' &nbsp;|&nbsp; '.join(spikes)}</div>""",
                        unsafe_allow_html=True)

        st.markdown(f"**Name:** {data.get('name','')}")
        ap = data.get("academic_profile", [])
        if ap:
            st.markdown(f"**Qualifications:** {len(ap)} entries")
            for e in ap:
                st.markdown(f"  • {e.get('degree','')} — {e.get('institution','')} — {e.get('score','')} ({e.get('year','')})")

        we = data.get("work_experience", [])
        if we:
            st.markdown(f"**Work/Internships:** {len(we)} entries")

        por = data.get("positions_of_responsibility", [])
        if por:
            st.markdown(f"**PORs:** {len(por)} entries")


# ── Main ─────────────────────────────────────────────────────────────────────

def main():
    init_state()
    sidebar()

    if st.session_state.stage == "upload":
        stage_upload()
    elif st.session_state.stage == "chat":
        stage_chat()
    elif st.session_state.stage == "download":
        stage_download()


if __name__ == "__main__":
    main()
