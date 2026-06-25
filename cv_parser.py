#!/usr/bin/env python3
"""
cv_parser.py — Extract raw text from uploaded CVs (PDF or DOCX)
Returns plain text that the agent then structures with Claude.
"""

import sys, os

def extract_pdf(path: str) -> str:
    try:
        import pdfplumber
        with pdfplumber.open(path) as pdf:
            pages = [p.extract_text() or '' for p in pdf.pages]
        return '\n'.join(pages)
    except Exception as e:
        return f"[PDF extraction error: {e}]"


def extract_docx(path: str) -> str:
    try:
        from docx import Document
        doc = Document(path)
        lines = []
        for para in doc.paragraphs:
            if para.text.strip():
                lines.append(para.text.strip())
        # Also extract table cells
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    lines.append('  |  '.join(cells))
        return '\n'.join(lines)
    except Exception as e:
        return f"[DOCX extraction error: {e}]"


def extract(path: str) -> str:
    """Auto-detect file type and extract text."""
    ext = os.path.splitext(path)[1].lower()
    if ext == '.pdf':
        return extract_pdf(path)
    elif ext in ('.docx', '.doc'):
        return extract_docx(path)
    elif ext in ('.txt', '.md'):
        with open(path, encoding='utf-8', errors='ignore') as f:
            return f.read()
    else:
        return f"[Unsupported file type: {ext}]"


if __name__ == '__main__':
    if len(sys.argv) < 2:
        print("Usage: python cv_parser.py <cv_file>")
        sys.exit(1)
    text = extract(sys.argv[1])
    print(text)
